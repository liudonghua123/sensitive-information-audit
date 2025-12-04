
from typing import Any, List
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.db import models
from app.schemas import Task, TaskCreate, Result, DatabaseMetadata, TableMetadata
from app.api import deps
from app.services.scanner import Scanner
from app.services.connector import DbConnector
import json
from fastapi.responses import StreamingResponse
from io import BytesIO

router = APIRouter()

@router.post("/", response_model=Task)
async def create_task(
    *,
    db: AsyncSession = Depends(deps.get_db),
    task_in: TaskCreate,
    background_tasks: BackgroundTasks,
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    # Serialize table_names to JSON if provided
    selected_tables_json = None
    if task_in.table_names:
        selected_tables_json = json.dumps(task_in.table_names)
    
    # Serialize rule_ids to JSON if provided
    selected_rules_json = None
    if task_in.rule_ids:
        selected_rules_json = json.dumps(task_in.rule_ids)
    
    task = models.ScanTask(
        connection_id=task_in.connection_id,
        status=models.TaskStatus.PENDING,
        selected_tables=selected_tables_json,
        selected_rules=selected_rules_json
    )
    db.add(task)
    await db.commit()
    await db.refresh(task)
    
    # Start background task
    scanner = Scanner(db)
    background_tasks.add_task(scanner.run_scan, task.id)
    
    return task

@router.get("/", response_model=List[Task])
async def read_tasks(
    db: AsyncSession = Depends(deps.get_db),
    skip: int = 0,
    limit: int = 100,
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    result = await db.execute(select(models.ScanTask).offset(skip).limit(limit))
    return result.scalars().all()

@router.get("/metadata/{connection_id}", response_model=DatabaseMetadata)
async def get_database_metadata(
    *,
    db: AsyncSession = Depends(deps.get_db),
    connection_id: int,
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    """Get metadata for a database connection including tables, views, columns, and statistics"""
    connection = await db.get(models.DbConnection, connection_id)
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    try:
        engine = DbConnector.create_engine(connection)
        metadata_list = await DbConnector.get_database_metadata(engine, connection.db_type)
        await engine.dispose()
        
        tables = [TableMetadata(**table_data) for table_data in metadata_list]
        return DatabaseMetadata(tables=tables)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch metadata: {str(e)}")

@router.get("/{id}/results", response_model=List[Result])
async def read_task_results(
    *,
    db: AsyncSession = Depends(deps.get_db),
    id: int,
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    result = await db.execute(select(models.ScanResult).where(models.ScanResult.task_id == id))
    return result.scalars().all()

@router.delete("/{id}", response_model=Any)
async def delete_task(
    *,
    db: AsyncSession = Depends(deps.get_db),
    id: int,
    current_user: models.User = Depends(deps.get_current_user),
) -> Any:
    result = await db.execute(select(models.ScanTask).where(models.ScanTask.id == id))
    task = result.scalars().first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    await db.delete(task)
    await db.commit()
    return {"ok": True}

@router.get("/{task_id}/export", response_class=StreamingResponse)
async def export_scan_report(
    task_id: int,
    format: str,
    db: AsyncSession = Depends(deps.get_db),
    current_user: models.User = Depends(deps.get_current_user),
):
    task_result = await db.execute(select(models.ScanTask).where(models.ScanTask.id == task_id))
    task = task_result.scalars().first()

    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    scan_results_query = await db.execute(select(models.ScanResult).where(models.ScanResult.task_id == task_id))
    scan_results = scan_results_query.scalars().all()

    if format.lower() == "docx":
        # Generate DOCX
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        document = Document()
        document.add_heading(f'Scan Report - Task #{task.id}', 0)
        
        # Task Info
        document.add_heading('Task Information', level=1)
        p = document.add_paragraph()
        p.add_run('Start Time: ').bold = True
        p.add_run(str(task.start_time) + '\n')
        p.add_run('End Time: ').bold = True
        p.add_run(str(task.end_time) + '\n')
        p.add_run('Status: ').bold = True
        p.add_run(task.status + '\n')
        
        # Summary
        if task.scan_summary:
            try:
                summary = json.loads(task.scan_summary)
                document.add_heading('Summary', level=1)
                
                table = document.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Value'
                
                row_cells = table.add_row().cells
                row_cells[0].text = 'Total Tables Scanned'
                row_cells[1].text = str(summary.get('total_tables_scanned', 0))
                
                row_cells = table.add_row().cells
                row_cells[0].text = 'Total Rows Scanned'
                row_cells[1].text = str(summary.get('total_rows_scanned', 0))
                
                row_cells = table.add_row().cells
                row_cells[0].text = 'Total Issues Found'
                row_cells[1].text = str(summary.get('total_issues_found', 0))
                
                # Rules Triggered
                if summary.get('rules_triggered'):
                    document.add_heading('Rules Triggered', level=2)
                    for rule, count in summary['rules_triggered'].items():
                        document.add_paragraph(f"{rule}: {count}", style='List Bullet')
                        
            except:
                pass
                
        # Detailed Results
        document.add_heading('Detailed Results', level=1)
        
        if not scan_results:
            document.add_paragraph('No issues found.')
        else:
            # Group by table
            grouped = {}
            for res in scan_results:
                if res.table_name not in grouped:
                    grouped[res.table_name] = []
                grouped[res.table_name].append(res)
                
            for table_name, results in grouped.items():
                document.add_heading(f'Table: {table_name}', level=2)
                
                table = document.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Column'
                hdr_cells[1].text = 'Rule'
                hdr_cells[2].text = 'Content'
                
                for res in results:
                    row_cells = table.add_row().cells
                    row_cells[0].text = res.column_name
                    row_cells[1].text = res.rule_name or ''
                    # Truncate content if too long
                    content = res.sensitive_content_masked or ''
                    if len(content) > 100:
                        content = content[:100] + '...'
                    row_cells[2].text = content

        # Save to buffer
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        headers = {
            'Content-Disposition': f'attachment; filename="scan_report_{task_id}.docx"'
        }
        return StreamingResponse(buffer, headers=headers, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    elif format.lower() == "pdf":
        # Generate PDF using ReportLab
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph(f"Scan Report - Task #{task.id}", styles['Title']))
        story.append(Spacer(1, 12))
        
        # Task Info
        story.append(Paragraph("Task Information", styles['Heading1']))
        story.append(Paragraph(f"<b>Start Time:</b> {task.start_time}", styles['Normal']))
        story.append(Paragraph(f"<b>End Time:</b> {task.end_time}", styles['Normal']))
        story.append(Paragraph(f"<b>Status:</b> {task.status}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Summary
        if task.scan_summary:
            try:
                summary = json.loads(task.scan_summary)
                story.append(Paragraph("Summary", styles['Heading1']))
                
                data = [
                    ['Metric', 'Value'],
                    ['Total Tables Scanned', str(summary.get('total_tables_scanned', 0))],
                    ['Total Rows Scanned', str(summary.get('total_rows_scanned', 0))],
                    ['Total Issues Found', str(summary.get('total_issues_found', 0))]
                ]
                
                t = Table(data, colWidths=[200, 100])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
                
                # Rules Triggered
                if summary.get('rules_triggered'):
                    story.append(Paragraph("Rules Triggered", styles['Heading2']))
                    for rule, count in summary['rules_triggered'].items():
                        story.append(Paragraph(f"• {rule}: {count}", styles['Normal']))
                    story.append(Spacer(1, 12))
            except:
                pass
        
        # Detailed Results
        story.append(Paragraph("Detailed Results", styles['Heading1']))
        
        if not scan_results:
            story.append(Paragraph("No issues found.", styles['Normal']))
        else:
            # Group by table
            grouped = {}
            for res in scan_results:
                if res.table_name not in grouped:
                    grouped[res.table_name] = []
                grouped[res.table_name].append(res)
                
            for table_name, results in grouped.items():
                story.append(Paragraph(f"Table: {table_name}", styles['Heading2']))
                
                data = [['Column', 'Rule', 'Content']]
                for res in results:
                    content = res.sensitive_content_masked or ''
                    if len(content) > 50:
                        content = content[:50] + '...'
                    data.append([
                        res.column_name,
                        res.rule_name or '',
                        content
                    ])
                
                t = Table(data, colWidths=[100, 100, 300])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        
        headers = {
            'Content-Disposition': f'attachment; filename="scan_report_{task_id}.pdf"'
        }
        return StreamingResponse(buffer, headers=headers, media_type='application/pdf')
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")